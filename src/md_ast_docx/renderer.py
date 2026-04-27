"""Render the internal model to a python-docx document."""

from __future__ import annotations

import csv
import io
import warnings
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path

from docx.document import Document as DocxDoc
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph as DocxParagraph

from md_ast_docx.errors import RenderError
from md_ast_docx.model import (
    Block,
    BlockQuote,
    BulletList,
    CodeBlock,
    CsvFileEmbed,
    CsvInlineEmbed,
    Document,
    Emphasis,
    Heading,
    ImageBlock,
    Inline,
    InlineCode,
    InlineImage,
    LineBreak,
    Link,
    ListItem,
    OrderedList,
    Paragraph,
    Strong,
    Table,
    TableCell,
    Text,
)
from md_ast_docx.options import ConversionOptions
from md_ast_docx.template import style_exists

_HYPERLINK_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/"
    "relationships/hyperlink"
)
_VALID_LINK_SCHEMES = ("http://", "https://", "mailto:")


@dataclass
class _RunFormat:
    bold: bool = False
    italic: bool = False
    code: bool = False


@dataclass
class _RunSegment:
    text: str
    fmt: _RunFormat = field(default_factory=_RunFormat)
    line_break_after: bool = False


class Renderer:
    """Render a parsed :class:`Document` into a DOCX document."""

    def __init__(
        self,
        doc: DocxDoc,
        options: ConversionOptions,
        markdown_dir: Path | None,
    ) -> None:
        self.doc = doc
        self.options = options
        self.markdown_dir = markdown_dir
        self.figure_counter = 0
        self.table_counter = 0

    # --- public entry -------------------------------------------------

    def render(self, document: Document) -> None:
        for block in document.blocks:
            self._render_block(block)

    # --- block dispatch -----------------------------------------------

    def _render_block(self, block: Block) -> None:
        if isinstance(block, Heading):
            self._render_heading(block)
        elif isinstance(block, Paragraph):
            self._render_paragraph(block)
        elif isinstance(block, CodeBlock):
            self._render_code_block(block)
        elif isinstance(block, BlockQuote):
            self._render_blockquote(block)
        elif isinstance(block, BulletList):
            self._render_list(block, ordered=False, level=1)
        elif isinstance(block, OrderedList):
            self._render_list(block, ordered=True, level=1)
        elif isinstance(block, Table):
            self._render_table(block)
        elif isinstance(block, ImageBlock):
            self._render_image_block(block)
        elif isinstance(block, CsvFileEmbed):
            self._render_csv_file(block)
        elif isinstance(block, CsvInlineEmbed):
            self._render_csv_inline(block)
        else:
            self._warn_or_raise(
                f"Unsupported block type: {type(block).__name__}"
            )

    # --- headings -----------------------------------------------------

    def _render_heading(self, h: Heading) -> None:
        style = self._heading_style(h.level)
        para = self.doc.add_paragraph(style=style)
        self._render_inlines(para, h.children)

    def _heading_style(self, level: int) -> str:
        for lvl in range(level, 0, -1):
            name = f"Heading {lvl}"
            if style_exists(self.doc, name):
                return name
        if level > 1:
            warnings.warn(
                f"No Heading style available for level {level}; "
                f"falling back to Normal",
                stacklevel=3,
            )
        return "Normal"

    # --- paragraph ----------------------------------------------------

    def _render_paragraph(self, p: Paragraph) -> None:
        para = self.doc.add_paragraph()
        deferred: list[InlineImage] = []
        self._render_inlines(para, p.children, deferred_images=deferred)
        for img in deferred:
            self._emit_figure(img.src, img.alt, after_paragraph=True)

    # --- code block ---------------------------------------------------

    def _render_code_block(self, cb: CodeBlock) -> None:
        if style_exists(self.doc, "Code"):
            para = self.doc.add_paragraph(style="Code")
            para.add_run(cb.text.rstrip("\n"))
        else:
            para = self.doc.add_paragraph()
            run = para.add_run(cb.text.rstrip("\n"))
            run.font.name = "Consolas"
            run.font.size = Pt(10)

    # --- blockquote ---------------------------------------------------

    def _render_blockquote(self, bq: BlockQuote) -> None:
        quote_style = "Quote" if style_exists(self.doc, "Quote") else "Normal"
        for inner in bq.blocks:
            if isinstance(inner, Paragraph):
                para = self.doc.add_paragraph(style=quote_style)
                self._render_inlines(para, inner.children)
            else:
                self._render_block(inner)

    # --- lists --------------------------------------------------------

    def _render_list(
        self,
        lst: BulletList | OrderedList,
        ordered: bool,
        level: int,
    ) -> None:
        if ordered and isinstance(lst, OrderedList) and lst.start != 1:
            warnings.warn(
                "Ordered list start value is not honored in v1 "
                "(Word numbering uses template defaults)",
                stacklevel=3,
            )
        for item in lst.items:
            self._render_list_item(item, ordered, level)

    def _render_list_item(
        self, item: ListItem, ordered: bool, level: int
    ) -> None:
        style = self._list_style(ordered, level)
        first_para_emitted = False
        for blk in item.blocks:
            if isinstance(blk, Paragraph) and not first_para_emitted:
                para = self.doc.add_paragraph(style=style)
                self._render_inlines(para, blk.children)
                first_para_emitted = True
            elif isinstance(blk, (BulletList, OrderedList)):
                self._render_list(
                    blk,
                    ordered=isinstance(blk, OrderedList),
                    level=level + 1,
                )
            else:
                self._render_block(blk)

    def _list_style(self, ordered: bool, level: int) -> str:
        base = "List Number" if ordered else "List Bullet"
        candidates = []
        if level > 1:
            candidates.append(f"{base} {level}")
        candidates.append(base)
        candidates.append("List Paragraph")
        candidates.append("Normal")
        for name in candidates:
            if style_exists(self.doc, name):
                return name
        return "Normal"

    # --- tables -------------------------------------------------------

    def _render_table(self, t: Table) -> None:
        n_cols = max(
            len(t.header.cells),
            *(len(r.cells) for r in t.body),
            1,
        )
        if t.caption is not None:
            self.table_counter += 1
            self._emit_caption(
                self.options.table_caption_prefix,
                self.table_counter,
                t.caption,
            )
        table_style = (
            "Table Grid" if style_exists(self.doc, "Table Grid") else None
        )
        n_rows = 1 + len(t.body)
        table = self.doc.add_table(rows=n_rows, cols=n_cols)
        if table_style:
            table.style = table_style
        self._fill_row(
            table.rows[0].cells, t.header.cells, t.alignments, bold=True
        )
        for i, row in enumerate(t.body, start=1):
            self._fill_row(
                table.rows[i].cells,
                row.cells,
                t.alignments,
                bold=False,
            )

    def _fill_row(
        self,
        docx_cells,
        model_cells: list[TableCell],
        alignments: list[str | None],
        bold: bool,
    ) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
        }
        for i, dcell in enumerate(docx_cells):
            mcell = model_cells[i] if i < len(model_cells) else TableCell()
            para = dcell.paragraphs[0]
            self._render_inlines(para, mcell.children)
            if bold:
                for run in para.runs:
                    run.bold = True
            align = alignments[i] if i < len(alignments) else None
            if align in align_map:
                para.alignment = align_map[align]

    # --- images / figure captions -------------------------------------

    def _render_image_block(self, ib: ImageBlock) -> None:
        self._emit_figure(ib.src, ib.alt, after_paragraph=False)

    def _emit_figure(self, src: str, alt: str, after_paragraph: bool) -> None:
        path = self._resolve_asset_path(src, kind="image")
        if path is None:
            return
        if not after_paragraph:
            para = self.doc.add_paragraph()
            run = para.add_run()
            try:
                run.add_picture(str(path))
            except Exception as exc:  # noqa: BLE001
                self._warn_or_raise(f"Failed to embed image {path}: {exc}")
                return
        # else: image was rendered inline already
        self.figure_counter += 1
        cap_inlines: list[Inline] = []
        if alt:
            cap_inlines.append(Text(alt))
        self._emit_caption(
            self.options.figure_caption_prefix,
            self.figure_counter,
            cap_inlines,
        )

    def _resolve_asset_path(
        self, src: str, kind: str = "asset"
    ) -> Path | None:
        if src.startswith(("http://", "https://")):
            self._warn_or_raise(f"Remote {kind} not supported: {src}")
            return None
        candidate = Path(src)
        resolved = (
            candidate
            if candidate.is_absolute()
            else (self._asset_base() / candidate).resolve()
        )
        if not resolved.exists():
            self._warn_or_raise(f"{kind.capitalize()} not found: {resolved}")
            return None
        return resolved

    def _asset_base(self) -> Path:
        if self.options.project_root is not None:
            return Path(self.options.project_root)
        if self.markdown_dir is not None:
            return self.markdown_dir
        return Path.cwd()

    # --- CSV embedding -----------------------------------------------

    def _render_csv_file(self, c: CsvFileEmbed) -> None:
        path = self._resolve_asset_path(c.path, kind="CSV file")
        if path is None:
            return
        try:
            text = path.read_text(encoding="utf-8")
        except (OSError, UnicodeDecodeError) as exc:
            self._warn_or_raise(f"Failed to read CSV {path}: {exc}")
            return
        rows = self._parse_csv_text(text)
        if not rows:
            self._warn_or_raise(f"CSV is empty: {path}")
            return
        self._emit_csv_table(rows, c.has_header, c.caption)

    def _render_csv_inline(self, c: CsvInlineEmbed) -> None:
        rows = self._parse_csv_text(c.data)
        if not rows:
            self._warn_or_raise("Inline CSV block is empty")
            return
        self._emit_csv_table(rows, c.has_header, c.caption)

    def _parse_csv_text(self, text: str) -> list[list[str]]:
        if not text.strip():
            return []
        try:
            dialect = csv.Sniffer().sniff(text[:1024], delimiters=",;\t|")
        except csv.Error:
            dialect = csv.excel
        return list(csv.reader(io.StringIO(text), dialect=dialect))

    def _emit_csv_table(
        self,
        rows: list[list[str]],
        has_header: bool,
        caption: list[Inline] | None,
    ) -> None:
        header = rows[0] if has_header else None
        body = rows[1:] if has_header else rows
        n_cols = max(
            len(header or []),
            *(len(r) for r in body),
            1,
        )
        if caption is not None:
            self.table_counter += 1
            self._emit_caption(
                self.options.table_caption_prefix,
                self.table_counter,
                caption,
            )
        n_rows = (1 if header is not None else 0) + len(body)
        if n_rows == 0:
            return
        table = self.doc.add_table(rows=n_rows, cols=n_cols)
        if style_exists(self.doc, "Table Grid"):
            table.style = "Table Grid"
        row_idx = 0
        if header is not None:
            for i, dcell in enumerate(table.rows[0].cells):
                cell_text = header[i] if i < len(header) else ""
                run = dcell.paragraphs[0].add_run(cell_text)
                run.bold = True
            row_idx = 1
        for body_row in body:
            for i, dcell in enumerate(table.rows[row_idx].cells):
                cell_text = body_row[i] if i < len(body_row) else ""
                dcell.paragraphs[0].add_run(cell_text)
            row_idx += 1

    # --- captions -----------------------------------------------------

    def _emit_caption(
        self, prefix: str, number: int, text_inlines: list[Inline]
    ) -> None:
        style = "Caption" if style_exists(self.doc, "Caption") else "Normal"
        para = self.doc.add_paragraph(style=style)
        if style == "Normal":
            # tasteful fallback formatting
            label = para.add_run(f"{prefix} ")
            label.italic = True
        else:
            para.add_run(f"{prefix} ")
        self._append_seq_field(para, prefix, str(number))
        if text_inlines:
            sep_run = para.add_run(": ")
            if style == "Normal":
                sep_run.italic = True
            self._render_inlines(
                para, text_inlines, force_italic=(style == "Normal")
            )

    def _append_seq_field(
        self, para: DocxParagraph, name: str, displayed: str
    ) -> None:
        """Append a Word SEQ field producing an arabic counter."""
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), f" SEQ {name} \\* ARABIC ")
        run = OxmlElement("w:r")
        text_el = OxmlElement("w:t")
        text_el.text = displayed
        run.append(text_el)
        fld.append(run)
        para._p.append(fld)

    # --- inlines ------------------------------------------------------

    def _render_inlines(
        self,
        para: DocxParagraph,
        inlines: list[Inline],
        deferred_images: list[InlineImage] | None = None,
        force_italic: bool = False,
    ) -> None:
        for inline in inlines:
            self._render_inline(
                para,
                inline,
                _RunFormat(italic=force_italic),
                deferred_images,
            )

    def _render_inline(
        self,
        para: DocxParagraph,
        inline: Inline,
        fmt: _RunFormat,
        deferred_images: list[InlineImage] | None,
    ) -> None:
        if isinstance(inline, Text):
            self._add_run(para, inline.text, fmt)
        elif isinstance(inline, LineBreak):
            run = para.add_run()
            run.add_break()
        elif isinstance(inline, InlineCode):
            new_fmt = _RunFormat(bold=fmt.bold, italic=fmt.italic, code=True)
            self._add_run(para, inline.text, new_fmt)
        elif isinstance(inline, Strong):
            new_fmt = _RunFormat(bold=True, italic=fmt.italic, code=fmt.code)
            for child in inline.children:
                self._render_inline(para, child, new_fmt, deferred_images)
        elif isinstance(inline, Emphasis):
            new_fmt = _RunFormat(bold=fmt.bold, italic=True, code=fmt.code)
            for child in inline.children:
                self._render_inline(para, child, new_fmt, deferred_images)
        elif isinstance(inline, Link):
            self._render_link(para, inline, fmt)
        elif isinstance(inline, InlineImage):
            self._render_inline_image(para, inline, deferred_images)
        else:
            self._warn_or_raise(f"Unsupported inline: {type(inline).__name__}")

    def _add_run(
        self, para: DocxParagraph, text: str, fmt: _RunFormat
    ) -> None:
        if not text:
            return
        run = para.add_run(text)
        if fmt.bold:
            run.bold = True
        if fmt.italic:
            run.italic = True
        if fmt.code:
            run.font.name = "Consolas"

    def _render_inline_image(
        self,
        para: DocxParagraph,
        img: InlineImage,
        deferred_images: list[InlineImage] | None,
    ) -> None:
        path = self._resolve_asset_path(img.src, kind="image")
        if path is None:
            return
        try:
            para.add_run().add_picture(str(path))
        except Exception as exc:  # noqa: BLE001
            self._warn_or_raise(f"Failed to embed image {path}: {exc}")
            return
        if deferred_images is not None:
            deferred_images.append(img)

    # --- hyperlinks ---------------------------------------------------

    def _render_link(
        self,
        para: DocxParagraph,
        link: Link,
        fmt: _RunFormat,
    ) -> None:
        href = link.href.strip()
        if not href.lower().startswith(_VALID_LINK_SCHEMES):
            # render as plain text
            warnings.warn(
                f"Skipping hyperlink with unsupported scheme: {href}",
                stacklevel=3,
            )
            for child in link.children:
                self._render_inline(para, child, fmt, None)
            return
        try:
            r_id = para.part.relate_to(href, _HYPERLINK_REL, is_external=True)
        except Exception as exc:  # noqa: BLE001
            warnings.warn(
                f"Hyperlink relationship failed for {href}: {exc}",
                stacklevel=3,
            )
            for child in link.children:
                self._render_inline(para, child, fmt, None)
            return
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        segments = self._collect_link_segments(link.children, fmt)
        for seg in segments:
            hyperlink.append(self._build_hyperlink_run(seg))
        para._p.append(hyperlink)

    def _collect_link_segments(
        self, inlines: list[Inline], base: _RunFormat
    ) -> list[_RunSegment]:
        segments: list[_RunSegment] = []
        self._walk_link_inlines(inlines, base, segments)
        return segments

    def _walk_link_inlines(
        self,
        inlines: list[Inline],
        fmt: _RunFormat,
        out: list[_RunSegment],
    ) -> None:
        for node in inlines:
            if isinstance(node, Text):
                if node.text:
                    out.append(_RunSegment(text=node.text, fmt=deepcopy(fmt)))
            elif isinstance(node, InlineCode):
                out.append(
                    _RunSegment(
                        text=node.text,
                        fmt=_RunFormat(
                            bold=fmt.bold,
                            italic=fmt.italic,
                            code=True,
                        ),
                    )
                )
            elif isinstance(node, Strong):
                self._walk_link_inlines(
                    node.children,
                    _RunFormat(
                        bold=True,
                        italic=fmt.italic,
                        code=fmt.code,
                    ),
                    out,
                )
            elif isinstance(node, Emphasis):
                self._walk_link_inlines(
                    node.children,
                    _RunFormat(
                        bold=fmt.bold,
                        italic=True,
                        code=fmt.code,
                    ),
                    out,
                )
            elif isinstance(node, LineBreak):
                if out:
                    out[-1].line_break_after = True

    def _build_hyperlink_run(self, seg: _RunSegment):
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        rPr.append(color)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        if seg.fmt.bold:
            rPr.append(OxmlElement("w:b"))
        if seg.fmt.italic:
            rPr.append(OxmlElement("w:i"))
        if seg.fmt.code:
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), "Consolas")
            rFonts.set(qn("w:hAnsi"), "Consolas")
            rPr.append(rFonts)
        run.append(rPr)
        t = OxmlElement("w:t")
        t.text = seg.text
        t.set(qn("xml:space"), "preserve")
        run.append(t)
        if seg.line_break_after:
            run.append(OxmlElement("w:br"))
        return run

    # --- helpers ------------------------------------------------------

    def _warn_or_raise(self, msg: str) -> None:
        if self.options.strict_mode:
            raise RenderError(msg)
        warnings.warn(msg, stacklevel=3)
