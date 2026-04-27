"""Tests for Jinja2 context substitution."""

from __future__ import annotations

import pytest
from docx import Document as DocxDocument
from docx.oxml.ns import qn

from md_reports import (
    ConversionOptions,
    MarkdownConverter,
    convert_markdown_text,
)
from md_reports.context import apply_context
from md_reports.errors import ValidationError
from md_reports.parser import parse

# --- apply_context unit tests ----------------------------------------


def test_no_context_is_noop():
    src = "hello {{ name }}"
    assert apply_context(src, None, strict=False) == src
    assert apply_context(src, {}, strict=False) == src


def test_simple_substitution():
    out = apply_context("hello {{ name }}", {"name": "Alice"}, False)
    assert out == "hello Alice"


def test_whitespace_variants():
    ctx = {"x": "v"}
    assert apply_context("{{x}}", ctx, False) == "v"
    assert apply_context("{{ x }}", ctx, False) == "v"
    assert apply_context("{{   x   }}", ctx, False) == "v"


def test_int_float_bool_values():
    out = apply_context(
        "{{ n }} {{ pi }} {{ ok }}",
        {"n": 5, "pi": 3.14, "ok": True},
        False,
    )
    assert out == "5 3.14 True"


def test_list_iteration():
    src = "{% for i in items %}- {{ i }}\n{% endfor %}"
    out = apply_context(src, {"items": [1, 2, 3]}, False)
    assert out == "- 1\n- 2\n- 3\n"


def test_dict_attribute_access():
    src = "{{ user.name }} ({{ user.role }})"
    out = apply_context(
        src, {"user": {"name": "Alice", "role": "Admin"}}, False
    )
    assert out == "Alice (Admin)"


def test_filter_works():
    out = apply_context("{{ name | upper }}", {"name": "alice"}, False)
    assert out == "ALICE"


def test_conditional_works():
    src = "{% if show %}YES{% else %}NO{% endif %}"
    assert apply_context(src, {"show": True}, False) == "YES"
    assert apply_context(src, {"show": False}, False) == "NO"


# --- missing variable behavior --------------------------------------


def test_missing_variable_non_strict_keeps_literal_and_warns():
    with pytest.warns(UserWarning):
        out = apply_context(
            "hello {{ missing }} world", {"other": 1}, strict=False
        )
    assert out == "hello {{ missing }} world"


def test_missing_variable_strict_raises():
    with pytest.raises(ValidationError):
        apply_context("hello {{ missing }}", {"other": 1}, strict=True)


def test_jinja_syntax_error_non_strict_warns_and_leaves_text():
    src = "broken {{ unclosed"
    with pytest.warns(UserWarning):
        out = apply_context(src, {"x": 1}, strict=False)
    assert out == src


def test_jinja_syntax_error_strict_raises():
    with pytest.raises(ValidationError):
        apply_context("broken {{ unclosed", {"x": 1}, strict=True)


# --- end-to-end through the parser/api ------------------------------


def _full_text(p) -> str:
    return "".join(t.text or "" for t in p._p.iter(qn("w:t")))


def test_substitution_in_heading_and_body(tmp_path):
    src = "# Q{{ q }} report\n\nRevenue grew {{ pct }}%."
    out = tmp_path / "out.docx"
    convert_markdown_text(src, out, context={"q": 2, "pct": 14.5})
    doc = DocxDocument(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "Q2 report" in texts
    assert any("Revenue grew 14.5%" in t for t in texts)


def test_substitution_in_table_cell(tmp_path):
    src = (
        "| Region | Revenue |\n"
        "|--------|---------|\n"
        "| EMEA   | {{ emea }} |\n"
        "| APAC   | {{ apac }} |\n"
    )
    out = tmp_path / "tbl.docx"
    convert_markdown_text(src, out, context={"emea": 10, "apac": 7})
    t = DocxDocument(str(out)).tables[0]
    assert t.cell(1, 1).text == "10"
    assert t.cell(2, 1).text == "7"


def test_substitution_in_csv_file_path(tmp_path):
    csv = tmp_path / "real_data.csv"
    csv.write_text("a,b\n1,2\n", encoding="utf-8")
    src = "```csv-file\n{{ data_file }}\n```\n"
    out = tmp_path / "csvpath.docx"
    convert_markdown_text(
        src,
        out,
        context={"data_file": "real_data.csv"},
        options=ConversionOptions(project_root=tmp_path),
    )
    t = DocxDocument(str(out)).tables[0]
    assert t.cell(0, 0).text == "a"


def test_substitution_in_inline_csv(tmp_path):
    src = "```csv\nregion,value\nEMEA,{{ emea }}\nAPAC,{{ apac }}\n```\n"
    out = tmp_path / "inline.docx"
    convert_markdown_text(src, out, context={"emea": 99, "apac": 12})
    t = DocxDocument(str(out)).tables[0]
    assert t.cell(1, 1).text == "99"
    assert t.cell(2, 1).text == "12"


def test_for_loop_emits_bullet_list(tmp_path):
    src = "Findings:\n\n{% for item in items %}- {{ item }}\n{% endfor %}\n"
    out = tmp_path / "loop.docx"
    convert_markdown_text(
        src, out, context={"items": ["alpha", "beta", "gamma"]}
    )
    doc = DocxDocument(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "alpha" in texts
    assert "beta" in texts
    assert "gamma" in texts


def test_conditional_block_skipped(tmp_path):
    src = "Always.\n\n{% if appendix %}Hidden appendix.{% endif %}\n"
    out = tmp_path / "cond.docx"
    convert_markdown_text(src, out, context={"appendix": False})
    doc = DocxDocument(str(out))
    texts = " ".join(p.text for p in doc.paragraphs)
    assert "Always." in texts
    assert "Hidden appendix" not in texts


def test_filter_in_markdown(tmp_path):
    src = "# {{ title | upper }}"
    out = tmp_path / "filter.docx"
    convert_markdown_text(src, out, context={"title": "hello"})
    doc = DocxDocument(str(out))
    assert any(p.text == "HELLO" for p in doc.paragraphs)


def test_strict_mode_missing_key_raises_via_api(tmp_path):
    out = tmp_path / "strict.docx"
    with pytest.raises(ValidationError):
        convert_markdown_text(
            "{{ missing }}",
            out,
            context={"other": 1},
            options=ConversionOptions(strict_mode=True),
        )


def test_non_strict_missing_key_keeps_literal_in_doc(tmp_path):
    out = tmp_path / "literal.docx"
    with pytest.warns(UserWarning):
        convert_markdown_text(
            "Value: {{ missing }} end",
            out,
            context={"other": 1},
        )
    doc = DocxDocument(str(out))
    texts = " ".join(p.text for p in doc.paragraphs)
    assert "{{ missing }}" in texts


# --- MarkdownConverter default_context --------------------------


def test_converter_default_context(tmp_path):
    conv = MarkdownConverter(default_context={"who": "World"})
    out = tmp_path / "a.docx"
    conv.convert_text("Hello {{ who }}", out)
    doc = DocxDocument(str(out))
    assert any("Hello World" == p.text for p in doc.paragraphs)


def test_converter_per_call_context_overrides_default(tmp_path):
    conv = MarkdownConverter(default_context={"who": "World"})
    out = tmp_path / "b.docx"
    conv.convert_text("Hello {{ who }}", out, context={"who": "Alice"})
    doc = DocxDocument(str(out))
    assert any("Hello Alice" == p.text for p in doc.paragraphs)


def test_converter_per_call_extends_default(tmp_path):
    conv = MarkdownConverter(default_context={"site": "Acme"})
    out = tmp_path / "c.docx"
    conv.convert_text("{{ site }} - {{ doc }}", out, context={"doc": "Q1"})
    doc = DocxDocument(str(out))
    assert any("Acme - Q1" == p.text for p in doc.paragraphs)


def test_parse_function_accepts_context():
    doc = parse("# {{ title }}", context={"title": "Hi"})
    heading = doc.blocks[0]
    # the heading's first inline child should be Text("Hi")
    from md_reports.model import Text

    assert isinstance(heading.children[0], Text)
    assert heading.children[0].text == "Hi"
