"""Tests for the security-hardening options.

Covers:
* ``ConversionOptions.sandboxed_context`` — Jinja2 SSTI sink hardened
  by switching to ``SandboxedEnvironment``.
* ``ConversionOptions.confine_assets`` — image/CSV path-traversal
  hardened by enforcing the resolved path stays under the asset base.
* Bookmark-name collision detection in cross-reference labels.
"""

from __future__ import annotations

import pytest
from docx import Document

from md_reports import (
    ConversionOptions,
    convert_markdown_file,
    convert_markdown_text,
)
from md_reports.errors import RenderError, ValidationError

# --- sandboxed Jinja2 -----------------------------------------------------


def test_sandboxed_context_blocks_attribute_escape(tmp_path):
    """Standard SSTI vector — accessing ``__class__`` on a context
    value — must fail under sandboxed mode."""
    md = "{{ name.__class__.__mro__ }}"
    out = tmp_path / "ssti.docx"
    with pytest.raises(ValidationError):
        convert_markdown_text(
            md,
            out,
            context={"name": "alice"},
            options=ConversionOptions(
                strict_mode=True,
                sandboxed_context=True,
            ),
        )


def test_sandboxed_context_warns_in_non_strict(tmp_path):
    md = "{{ name.__class__ }}"
    out = tmp_path / "ssti2.docx"
    with pytest.warns(UserWarning):
        convert_markdown_text(
            md,
            out,
            context={"name": "alice"},
            options=ConversionOptions(sandboxed_context=True),
        )
    # File still produced (markdown left unchanged on substitution
    # failure in non-strict mode)
    assert out.exists()


def test_sandboxed_context_allows_normal_substitution(tmp_path):
    """Sandboxing must not break ordinary variable substitution."""
    md = "# {{ title }}\n\nValue: **{{ pct }}**"
    out = tmp_path / "ok.docx"
    convert_markdown_text(
        md,
        out,
        context={"title": "Q4", "pct": 14.5},
        options=ConversionOptions(sandboxed_context=True),
    )
    doc = Document(str(out))
    paragraphs = [p.text for p in doc.paragraphs]
    assert "Q4" in paragraphs
    assert any("14.5" in p for p in paragraphs)


def test_unsandboxed_context_remains_default(tmp_path):
    """Backwards compatibility: developer-authored markdown can still
    introspect context values when sandboxing is off (default)."""
    md = "{{ name.upper() }}"
    out = tmp_path / "loud.docx"
    convert_markdown_text(
        md,
        out,
        context={"name": "alice"},
    )
    doc = Document(str(out))
    assert any("ALICE" in p.text for p in doc.paragraphs)


# --- confine_assets -------------------------------------------------------


def test_confine_assets_blocks_absolute_path_outside_base(tmp_path):
    """An absolute path pointing outside the asset base is rejected."""
    secret = tmp_path / "secret.csv"
    secret.write_text("a,b\n1,2\n", encoding="utf-8")

    base = tmp_path / "project"
    base.mkdir()
    md_path = base / "doc.md"
    abs_secret = str(secret).replace("\\", "/")
    md_path.write_text(
        f"```csv-file\n{abs_secret}\n```\n",
        encoding="utf-8",
    )

    out = tmp_path / "out.docx"
    with pytest.raises(RenderError, match="escapes asset base"):
        convert_markdown_file(
            md_path,
            out,
            options=ConversionOptions(
                strict_mode=True,
                confine_assets=True,
                project_root=base,
            ),
        )


def test_confine_assets_blocks_dotdot_traversal(tmp_path):
    secret = tmp_path / "secret.csv"
    secret.write_text("a,b\n1,2\n", encoding="utf-8")

    base = tmp_path / "project"
    base.mkdir()

    md = "```csv-file\n../secret.csv\n```\n"
    out = tmp_path / "out.docx"
    with pytest.raises(RenderError, match="escapes asset base"):
        convert_markdown_text(
            md,
            out,
            options=ConversionOptions(
                strict_mode=True,
                confine_assets=True,
                project_root=base,
            ),
        )


def test_confine_assets_allows_inbase_relative_path(tmp_path):
    base = tmp_path / "project"
    base.mkdir()
    csv = base / "data.csv"
    csv.write_text("a,b\n1,2\n", encoding="utf-8")

    md = "```csv-file\ndata.csv\n```\n"
    out = tmp_path / "out.docx"
    convert_markdown_text(
        md,
        out,
        options=ConversionOptions(
            strict_mode=True,
            confine_assets=True,
            project_root=base,
        ),
    )
    assert out.exists()


def test_confine_assets_off_by_default(tmp_path):
    """Default (off) preserves existing behaviour: absolute paths
    inside the test temp dir are accepted."""
    secret = tmp_path / "data.csv"
    secret.write_text("a,b\n1,2\n", encoding="utf-8")

    base = tmp_path / "project"
    base.mkdir()
    abs_path = str(secret).replace("\\", "/")
    md = f"```csv-file\n{abs_path}\n```\n"
    out = tmp_path / "out.docx"
    convert_markdown_text(
        md,
        out,
        options=ConversionOptions(project_root=base),
    )
    assert out.exists()


# --- bookmark collision ---------------------------------------------------


def test_bookmark_name_collision_warns(tmp_path):
    """Two distinct user labels that collapse to the same bookmark
    name after sanitisation must warn and keep only the first."""
    md = (
        "Table: First {#tab-x}\n\n| a |\n|---|\n| 1 |\n\n"
        "Table: Second {#tab_x}\n\n| a |\n|---|\n| 1 |\n"
    )
    out = tmp_path / "collide.docx"
    with pytest.warns(UserWarning, match="collides with existing label"):
        convert_markdown_text(md, out)


def test_bookmark_name_collision_raises_in_strict_mode(tmp_path):
    md = (
        "Table: First {#tab-x}\n\n| a |\n|---|\n| 1 |\n\n"
        "Table: Second {#tab_x}\n\n| a |\n|---|\n| 1 |\n"
    )
    out = tmp_path / "collide.docx"
    with pytest.raises(RenderError, match="collides with existing label"):
        convert_markdown_text(
            md,
            out,
            options=ConversionOptions(strict_mode=True),
        )
