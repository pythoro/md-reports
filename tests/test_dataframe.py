"""Tests for the ``to_csv`` Jinja2 filter / DataFrame embedding."""

from __future__ import annotations

import pytest
from docx import Document as DocxDocument
from docx.oxml.ns import qn

from md_reports import ConversionOptions, convert_markdown_text
from md_reports.context import apply_context


class _PandasLikeDF:
    """Minimal pandas-DataFrame stand-in.

    Mirrors enough of ``pandas.DataFrame.to_csv`` for our filter:
    accepts ``index=`` and ``sep=`` kwargs, returns a string with a
    trailing newline (matching pandas behavior).
    """

    def __init__(self, columns: list[str], rows: list[list[object]]) -> None:
        self.columns = columns
        self.rows = rows

    def to_csv(
        self,
        path_or_buf: object = None,
        *,
        index: bool = True,
        sep: str = ",",
    ) -> str:
        out: list[str] = []
        if index:
            out.append(sep.join([""] + list(self.columns)))
            for i, row in enumerate(self.rows):
                out.append(sep.join([str(i)] + [str(c) for c in row]))
        else:
            out.append(sep.join(self.columns))
            for row in self.rows:
                out.append(sep.join(str(c) for c in row))
        return "\n".join(out) + "\n"


class _NoIndexDF:
    """DataFrame-like whose to_csv() takes no ``index`` kwarg.

    Stands in for libraries (e.g., polars-style) where the default
    output already excludes any index. Verifies that the filter does
    not blindly inject ``index=False`` when the method's signature
    doesn't accept it.
    """

    def to_csv(self, *, sep: str = ",") -> str:
        return f"a{sep}b\n1{sep}2\n"


# --- filter unit tests ----------------------------------------------


def test_filter_basic_substitution():
    df = _PandasLikeDF(["a", "b"], [[1, 2], [3, 4]])
    out = apply_context("{{ df | to_csv }}", {"df": df}, strict=False)
    # index=False default — no leading index column
    assert out.strip().splitlines() == ["a,b", "1,2", "3,4"]


def test_filter_drops_trailing_newline():
    df = _PandasLikeDF(["x"], [[1]])
    out = apply_context("X{{ df | to_csv }}Y", {"df": df}, False)
    # No trailing blank line bleed-through
    assert out.startswith("Xx\n1Y") or out == "Xx\n1Y"


def test_filter_index_true_passthrough():
    df = _PandasLikeDF(["a", "b"], [[1, 2]])
    out = apply_context("{{ df | to_csv(index=True) }}", {"df": df}, False)
    lines = out.strip().splitlines()
    assert lines[0] == ",a,b"
    assert lines[1] == "0,1,2"


def test_filter_sep_passthrough():
    df = _PandasLikeDF(["a", "b"], [[1, 2]])
    out = apply_context("{{ df | to_csv(sep=';') }}", {"df": df}, False)
    assert out.strip().splitlines() == ["a;b", "1;2"]


def test_filter_skips_index_when_method_does_not_accept_it():
    df = _NoIndexDF()
    # Should not raise even though index=False isn't accepted
    out = apply_context("{{ df | to_csv }}", {"df": df}, False)
    assert out.strip().splitlines() == ["a,b", "1,2"]


def test_filter_rejects_non_dataframe():
    from md_reports.errors import ValidationError

    # Strict mode bubbles the underlying TypeError as ValidationError
    # via the apply_context error path.
    with pytest.raises(ValidationError):
        apply_context(
            "{{ s | to_csv }}", {"s": "not a dataframe"}, strict=True
        )


def test_filter_non_strict_warns_on_non_dataframe():
    src = "{{ s | to_csv }}"
    with pytest.warns(UserWarning):
        out = apply_context(src, {"s": "no method"}, strict=False)
    # Falls back to leaving text unchanged
    assert out == src


# --- end-to-end through the API ------------------------------------


def _full_text(p) -> str:
    return "".join(t.text or "" for t in p._p.iter(qn("w:t")))


def test_dataframe_renders_as_table_via_csv_fence(tmp_path):
    df = _PandasLikeDF(
        ["region", "q1", "q2"],
        [["EMEA", 1, 2], ["APAC", 3, 4]],
    )
    src = "```csv\n{{ df | to_csv }}\n```\n"
    out = tmp_path / "df.docx"
    convert_markdown_text(src, out, context={"df": df})
    doc = DocxDocument(str(out))
    assert len(doc.tables) == 1
    t = doc.tables[0]
    assert len(t.rows) == 3 and len(t.columns) == 3
    assert t.cell(0, 0).text == "region"
    assert t.cell(2, 2).text == "4"
    # header should be bolded
    assert t.cell(0, 0).paragraphs[0].runs[0].bold is True


def test_dataframe_with_caption(tmp_path):
    df = _PandasLikeDF(["a"], [[1], [2]])
    src = "Table: My data.\n\n```csv\n{{ df | to_csv }}\n```\n"
    out = tmp_path / "cap.docx"
    convert_markdown_text(src, out, context={"df": df})
    doc = DocxDocument(str(out))
    captions = [
        _full_text(p) for p in doc.paragraphs if p.style.name == "Caption"
    ]
    joined = " ".join(captions)
    assert "Table 1" in joined
    assert "My data" in joined


def test_dataframe_with_no_header_flag(tmp_path):
    df = _PandasLikeDF(["a", "b"], [[1, 2], [3, 4]])
    src = "```csv no-header\n{{ df | to_csv }}\n```\n"
    out = tmp_path / "nh.docx"
    convert_markdown_text(src, out, context={"df": df})
    t = DocxDocument(str(out)).tables[0]
    # With no-header, the column-name row becomes data — three rows
    # total, none bolded.
    assert len(t.rows) == 3
    assert not t.cell(0, 0).paragraphs[0].runs[0].bold
    assert t.cell(0, 0).text == "a"


def test_dataframe_shares_table_counter(tmp_path):
    df = _PandasLikeDF(["x"], [[1]])
    src = (
        "Table: First.\n\n"
        "| a |\n|---|\n| 1 |\n\n"
        "Table: Second.\n\n"
        "```csv\n{{ df | to_csv }}\n```\n"
    )
    out = tmp_path / "shared.docx"
    convert_markdown_text(src, out, context={"df": df})
    doc = DocxDocument(str(out))
    captions = [
        _full_text(p) for p in doc.paragraphs if p.style.name == "Caption"
    ]
    joined = " ".join(captions)
    assert "Table 1" in joined and "Table 2" in joined


def test_dataframe_strict_missing_method_raises(tmp_path):
    from md_reports.errors import ValidationError

    out = tmp_path / "bad.docx"
    with pytest.raises(ValidationError):
        convert_markdown_text(
            "```csv\n{{ s | to_csv }}\n```",
            out,
            context={"s": 42},
            options=ConversionOptions(strict_mode=True),
        )


def test_dataframe_with_sep_kwarg_renders_correctly(tmp_path):
    df = _PandasLikeDF(["a", "b"], [[1, 2]])
    src = "```csv\n{{ df | to_csv(sep=';') }}\n```\n"
    out = tmp_path / "sep.docx"
    convert_markdown_text(src, out, context={"df": df})
    t = DocxDocument(str(out)).tables[0]
    # Sniffer should pick up semicolon as the delimiter
    assert t.cell(0, 0).text == "a"
    assert t.cell(0, 1).text == "b"
    assert t.cell(1, 1).text == "2"
