"""CSV embedding tests (parser + renderer)."""

from __future__ import annotations

import pytest
from docx import Document as DocxDocument
from docx.oxml.ns import qn

from md_ast_docx import (
    ConversionOptions,
    convert_markdown_text,
)
from md_ast_docx.errors import RenderError
from md_ast_docx.model import (
    CodeBlock,
    CsvFileEmbed,
    CsvInlineEmbed,
    Table,
    Text,
)
from md_ast_docx.parser import parse

# --- parser -----------------------------------------------------------


def test_csv_file_fence_emits_csv_file_embed():
    src = "```csv-file\ndata/x.csv\n```\n"
    doc = parse(src)
    blk = doc.blocks[0]
    assert isinstance(blk, CsvFileEmbed)
    assert blk.path == "data/x.csv"
    assert blk.has_header is True


def test_csv_inline_fence_emits_inline_embed():
    src = "```csv\na,b\n1,2\n```\n"
    doc = parse(src)
    blk = doc.blocks[0]
    assert isinstance(blk, CsvInlineEmbed)
    assert "a,b" in blk.data and "1,2" in blk.data
    assert blk.has_header is True


def test_no_header_flag_parsed():
    for fence in (
        "```csv-file no-header\nx.csv\n```\n",
        "```csv no-header\n1,2,3\n```\n",
    ):
        blk = parse(fence).blocks[0]
        assert blk.has_header is False


def test_other_fences_still_codeblocks():
    blk = parse("```python\nprint(1)\n```\n").blocks[0]
    assert isinstance(blk, CodeBlock)
    assert blk.language == "python"


def test_table_caption_attaches_to_csv_file_embed():
    src = "Table: Quarterly.\n\n```csv-file\ndata/q.csv\n```\n"
    blk = parse(src).blocks[0]
    assert isinstance(blk, CsvFileEmbed)
    assert blk.caption is not None
    text = "".join(c.text for c in blk.caption if isinstance(c, Text))
    assert "Quarterly" in text


def test_table_caption_attaches_to_csv_inline_embed():
    src = "Table: Inline data.\n\n```csv\na,b\n1,2\n```\n"
    blk = parse(src).blocks[0]
    assert isinstance(blk, CsvInlineEmbed)
    assert blk.caption is not None


# --- renderer / integration ------------------------------------------


def _full_text(p) -> str:
    return "".join(t.text or "" for t in p._p.iter(qn("w:t")))


def test_csv_file_embedding_renders_table(tmp_path):
    csv = tmp_path / "data.csv"
    csv.write_text("region,q1,q2\nEMEA,1,2\nAPAC,3,4\n", encoding="utf-8")
    md = "# Title\n\nTable: Sales.\n\n```csv-file\ndata.csv\n```\n"
    out = tmp_path / "out.docx"
    convert_markdown_text(
        md, out, options=ConversionOptions(project_root=tmp_path)
    )
    doc = DocxDocument(str(out))
    assert len(doc.tables) == 1
    t = doc.tables[0]
    assert len(t.rows) == 3 and len(t.columns) == 3
    assert t.cell(0, 0).text == "region"
    assert t.cell(2, 2).text == "4"
    # header row should be bolded
    header_run = t.cell(0, 0).paragraphs[0].runs[0]
    assert header_run.bold is True
    # caption above the table
    body = list(doc.element.body)
    tbl_idx = next(i for i, el in enumerate(body) if el.tag == qn("w:tbl"))
    cap = body[tbl_idx - 1]
    assert "Sales" in "".join(cap.itertext())


def test_csv_inline_embedding_renders_table(tmp_path):
    md = "```csv\nregion,q1\nEMEA,1\nAPAC,3\n```\n"
    out = tmp_path / "inline.docx"
    convert_markdown_text(md, out)
    doc = DocxDocument(str(out))
    assert len(doc.tables) == 1
    t = doc.tables[0]
    assert t.cell(0, 0).text == "region"
    assert t.cell(2, 1).text == "3"


def test_no_header_flag_skips_bold_header(tmp_path):
    md = "```csv no-header\n1,2\n3,4\n```\n"
    out = tmp_path / "nh.docx"
    convert_markdown_text(md, out)
    t = DocxDocument(str(out)).tables[0]
    assert len(t.rows) == 2
    first_run = t.cell(0, 0).paragraphs[0].runs[0]
    # without a header the first row is just body — not bolded
    assert not first_run.bold


def test_missing_csv_file_warns_and_continues(tmp_path):
    md = "```csv-file\nnope.csv\n```\n"
    out = tmp_path / "miss.docx"
    with pytest.warns(UserWarning):
        convert_markdown_text(
            md, out, options=ConversionOptions(project_root=tmp_path)
        )
    assert out.exists()
    # no table emitted
    assert len(DocxDocument(str(out)).tables) == 0


def test_strict_missing_csv_raises(tmp_path):
    md = "```csv-file\nnope.csv\n```\n"
    out = tmp_path / "strict.docx"
    with pytest.raises(RenderError):
        convert_markdown_text(
            md,
            out,
            options=ConversionOptions(strict_mode=True, project_root=tmp_path),
        )


def test_shared_table_counter_with_markdown_tables(tmp_path):
    csv = tmp_path / "d.csv"
    csv.write_text("a,b\n1,2\n", encoding="utf-8")
    md = (
        "Table: First.\n\n"
        "| a | b |\n|---|---|\n| 1 | 2 |\n\n"
        "Table: Second.\n\n"
        "```csv-file\nd.csv\n```\n\n"
        "Table: Third.\n\n"
        "```csv\nx,y\n9,8\n```\n"
    )
    out = tmp_path / "mix.docx"
    convert_markdown_text(
        md, out, options=ConversionOptions(project_root=tmp_path)
    )
    doc = DocxDocument(str(out))
    captions = [
        _full_text(p) for p in doc.paragraphs if p.style.name == "Caption"
    ]
    joined = " ".join(captions)
    assert "Table 1" in joined
    assert "Table 2" in joined
    assert "Table 3" in joined


def test_semicolon_delimited_csv_sniffed(tmp_path):
    csv = tmp_path / "semi.csv"
    csv.write_text("a;b;c\n1;2;3\n", encoding="utf-8")
    md = "```csv-file\nsemi.csv\n```\n"
    out = tmp_path / "semi.docx"
    convert_markdown_text(
        md, out, options=ConversionOptions(project_root=tmp_path)
    )
    t = DocxDocument(str(out)).tables[0]
    assert t.cell(0, 0).text == "a"
    assert t.cell(0, 2).text == "c"
    assert t.cell(1, 1).text == "2"


def test_project_root_overrides_markdown_dir(tmp_path):
    # CSV located in tmp_path; markdown file in a sibling dir.
    sub = tmp_path / "docs"
    sub.mkdir()
    md_path = sub / "doc.md"
    csv = tmp_path / "data.csv"
    csv.write_text("h\n1\n", encoding="utf-8")
    md_path.write_text("```csv-file\ndata.csv\n```\n", encoding="utf-8")
    out = tmp_path / "rooted.docx"
    from md_ast_docx import convert_markdown_file

    convert_markdown_file(
        md_path,
        out,
        options=ConversionOptions(project_root=tmp_path),
    )
    t = DocxDocument(str(out)).tables[0]
    assert t.cell(0, 0).text == "h"


def test_csv_does_not_disturb_markdown_table_after_it(tmp_path):
    md = "```csv\na,b\n1,2\n```\n\n| x | y |\n|---|---|\n| 9 | 8 |\n"
    out = tmp_path / "two.docx"
    convert_markdown_text(md, out)
    doc = DocxDocument(str(out))
    assert len(doc.tables) == 2
    assert doc.tables[1].cell(1, 0).text == "9"


def test_imports_table_class_compatible():
    # Sanity check that Table model is still importable with shared
    # counter logic (no symbol drift in this refactor).
    assert Table is not None
