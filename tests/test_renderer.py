"""Renderer / integration tests."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.oxml.ns import qn

from md_reports import (
    ConversionOptions,
    MarkdownConverter,
    convert_markdown_file,
    convert_markdown_text,
    get_default_template_path,
)


def _open(path: Path):
    return DocxDocument(str(path))


def _full_paragraph_text(p) -> str:
    """Return all w:t text descendants of a paragraph."""
    return "".join(t.text or "" for t in p._p.iter(qn("w:t")))


def test_end_to_end_basic(tmp_path):
    out = tmp_path / "out.docx"
    convert_markdown_text("# Title\n\nHello **world**.\n", out)
    assert out.exists()
    doc = _open(out)
    paragraphs = [p.text for p in doc.paragraphs]
    assert "Title" in paragraphs
    assert any("Hello" in p and "world" in p for p in paragraphs)


def test_default_template_used_when_none_given(tmp_path):
    out = tmp_path / "default.docx"
    convert_markdown_text("plain text", out)
    assert out.exists()


def test_get_default_template_path_returns_existing_file():
    p = get_default_template_path()
    assert p.exists()


def test_convert_file_resolves_relative_image(tmp_path):
    # write a 1x1 PNG as an image to embed
    png = tmp_path / "tiny.png"
    png.write_bytes(_one_pixel_png())
    md = tmp_path / "doc.md"
    md.write_text(
        "# Title\n\n![A tiny picture](tiny.png)\n",
        encoding="utf-8",
    )
    out = tmp_path / "out.docx"
    convert_markdown_file(md, out)
    doc = _open(out)
    # Caption paragraph should appear
    caption_texts = [
        p.text for p in doc.paragraphs if p.style.name == "Caption"
    ]
    joined = " ".join(caption_texts)
    assert "Figure" in joined and "A tiny picture" in joined


def test_table_caption_above_table(tmp_path):
    md_text = "Table: My table.\n\n| a | b |\n|---|---|\n| 1 | 2 |\n"
    out = tmp_path / "tab.docx"
    convert_markdown_text(md_text, out)
    doc = _open(out)
    body_elems = list(doc.element.body)
    # find the first w:tbl
    tbl_idx = next(
        i for i, el in enumerate(body_elems) if el.tag == qn("w:tbl")
    )
    # at least one paragraph must precede the table
    assert tbl_idx >= 1
    prev = body_elems[tbl_idx - 1]
    assert prev.tag == qn("w:p")
    text = "".join(prev.itertext())
    assert "Table" in text and "My table" in text


def test_figure_and_table_counters_independent(tmp_path):
    png = tmp_path / "p.png"
    png.write_bytes(_one_pixel_png())
    md_text = (
        "![first](p.png)\n\n"
        "![second](p.png)\n\n"
        "Table: T1.\n\n"
        "| a |\n|---|\n| 1 |\n\n"
        "Table: T2.\n\n"
        "| a |\n|---|\n| 1 |\n"
    )
    out = tmp_path / "fig_tab.docx"
    convert_markdown_text(
        md_text,
        out,
        options=ConversionOptions(project_root=tmp_path),
    )
    doc = _open(out)
    captions = [
        _full_paragraph_text(p)
        for p in doc.paragraphs
        if p.style.name == "Caption"
    ]
    joined = " ".join(captions)
    assert "Figure 1" in joined
    assert "Figure 2" in joined
    assert "Table 1" in joined
    assert "Table 2" in joined


def test_hyperlink_relationship_added(tmp_path):
    out = tmp_path / "link.docx"
    convert_markdown_text("see [docs](https://example.com)", out)
    doc = _open(out)
    rels = doc.part.rels
    targets = {r.target_ref for r in rels.values()}
    assert "https://example.com" in targets


def test_html_anchor_renders_as_hyperlink(tmp_path):
    out = tmp_path / "html.docx"
    convert_markdown_text('visit <a href="https://x.test">x</a>', out)
    doc = _open(out)
    rels = doc.part.rels
    targets = {r.target_ref for r in rels.values()}
    assert "https://x.test" in targets


def test_unsupported_link_scheme_falls_back_to_text(tmp_path):
    # FTP is not in our allowed scheme list; renderer emits as plain
    # text and warns. (markdown-it strips `javascript:` URLs entirely
    # before they reach the renderer, so we don't use those here.)
    out = tmp_path / "ftp.docx"
    with pytest.warns(UserWarning):
        convert_markdown_text("[file](ftp://x.test/f)", out)
    doc = _open(out)
    rels = doc.part.rels
    targets = {r.target_ref for r in rels.values()}
    assert "ftp://x.test/f" not in targets


def test_missing_image_warns_then_continues(tmp_path):
    out = tmp_path / "missing.docx"
    with pytest.warns(UserWarning):
        convert_markdown_text(
            "![alt](does_not_exist.png)\n\nstill writes",
            out,
            options=ConversionOptions(project_root=tmp_path),
        )
    assert out.exists()


def test_strict_missing_image_raises(tmp_path):
    from md_reports.errors import RenderError

    out = tmp_path / "strict.docx"
    with pytest.raises(RenderError):
        convert_markdown_text(
            "![alt](nope.png)",
            out,
            options=ConversionOptions(strict_mode=True, project_root=tmp_path),
        )


def test_converter_class_reuses_template(tmp_path):
    conv = MarkdownConverter()
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    conv.convert_text("# A", a)
    conv.convert_text("# B", b)
    assert a.exists() and b.exists()


def test_table_renders_with_correct_dimensions(tmp_path):
    md_text = "| h1 | h2 |\n|---|---|\n| 1 | 2 |\n| 3 | 4 |\n"
    out = tmp_path / "tbl.docx"
    convert_markdown_text(md_text, out)
    doc = _open(out)
    assert len(doc.tables) == 1
    t = doc.tables[0]
    assert len(t.rows) == 3  # header + 2 body
    assert len(t.columns) == 2
    assert t.cell(0, 0).text == "h1"
    assert t.cell(2, 1).text == "4"


def test_heading_levels_h1_h6(tmp_path):
    src = "\n\n".join(f"{'#' * lvl} h{lvl}" for lvl in range(1, 7))
    out = tmp_path / "headings.docx"
    convert_markdown_text(src, out)
    doc = _open(out)
    style_by_text = {p.text: p.style.name for p in doc.paragraphs}
    for lvl in range(1, 7):
        assert style_by_text[f"h{lvl}"] == f"Heading {lvl}"


def test_seq_field_present_in_caption(tmp_path):
    md_text = "Table: cap.\n\n| a |\n|---|\n| 1 |\n"
    out = tmp_path / "seq.docx"
    convert_markdown_text(md_text, out)
    doc = _open(out)
    cap = next(p for p in doc.paragraphs if p.style.name == "Caption")
    fld_chars = cap._p.findall(f".//{qn('w:fldChar')}")
    types = [fc.get(qn("w:fldCharType")) for fc in fld_chars]
    assert "begin" in types and "separate" in types and "end" in types
    instr_texts = cap._p.findall(f".//{qn('w:instrText')}")
    assert instr_texts, "expected SEQ instrText in caption"
    joined = "".join(t.text or "" for t in instr_texts)
    assert "SEQ" in joined and "Table" in joined


def test_caption_uses_concrete_style_without_inline_italic(tmp_path):
    md_text = "Table: cap.\n\n| a |\n|---|\n| 1 |\n"
    out = tmp_path / "concrete.docx"
    convert_markdown_text(md_text, out)
    doc = _open(out)
    cap = next(p for p in doc.paragraphs if p.style.name == "Caption")
    # The default template's Caption style supplies italic, so the
    # renderer should not also stamp italic on the runs themselves.
    assert all(r.italic is None for r in cap.runs)


def test_caption_falls_back_to_inline_italic_when_style_missing(tmp_path):
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    from md_reports import DocxRenderer

    template_src = get_default_template_path()
    stripped = tmp_path / "stripped.docx"
    stripped.write_bytes(template_src.read_bytes())
    doc = DocxDocument(str(stripped))
    styles_el = doc.styles.element
    for s in list(styles_el.findall(qn("w:style"))):
        if s.get(qn("w:styleId")) == "Caption":
            styles_el.remove(s)
    doc.save(str(stripped))

    out = tmp_path / "fallback.docx"
    md_text = "Table: cap.\n\n| a |\n|---|\n| 1 |\n"
    convert_markdown_text(
        md_text,
        out,
        renderer=DocxRenderer(template_path=stripped),
    )
    doc2 = _open(out)
    cap = next(
        p for p in doc2.paragraphs if "Table" in _full_paragraph_text(p)
    )
    assert cap.style.name == "Normal"
    assert any(r.italic is True for r in cap.runs)


def _one_pixel_png() -> bytes:
    """A minimal valid 1x1 PNG."""
    import base64

    return base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQ"
        "VR42mNkYAAAAAYAAjCB0C8AAAAASUVORK5CYII="
    )
